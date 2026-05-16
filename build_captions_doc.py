from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── MoM brand colors ──────────────────────────────────────────
RED   = RGBColor(0xEC, 0x1C, 0x2D)
DARK  = RGBColor(0x1A, 0x1A, 0x1A)
BODY  = RGBColor(0x33, 0x2E, 0x2A)
MID   = RGBColor(0x6E, 0x66, 0x5B)
LIGHT = RGBColor(0xA0, 0x98, 0x90)

# Platform accent colors
PLATFORM_COLORS = {
    'Instagram': ('833AB4', RGBColor(0x83, 0x3A, 0xB4)),  # purple
    'Facebook':  ('1877F2', RGBColor(0x18, 0x77, 0xF2)),  # blue
    'LinkedIn':  ('0A66C2', RGBColor(0x0A, 0x66, 0xC2)),  # LinkedIn blue
    'X':         ('000000', RGBColor(0x00, 0x00, 0x00)),  # black
}

# ── Helpers ───────────────────────────────────────────────────
def set_cell_bg(cell, hex6):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6)
    tcPr.append(shd)

def set_top_border(cell, color_hex, sz=24):
    tcPr = cell._tc.get_or_add_tcPr()
    tcB  = OxmlElement('w:tcBorders')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    str(sz))
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), color_hex)
    tcB.append(top)
    for side in ('left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'E8DECE')
        tcB.append(el)
    tcPr.append(tcB)

def run(para, text, bold=False, italic=False, size=None, color=None):
    r = para.add_run(text)
    r.bold   = bold
    r.italic = italic
    if size:  r.font.size      = Pt(size)
    if color: r.font.color.rgb = color
    return r


# ── Caption card ──────────────────────────────────────────────
def caption_card(doc, number, label, text, platform_hex, platform_rgb, tip=None):
    """One caption in a framed card with platform-colored top border."""

    # small label above the box
    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after  = Pt(3)
    run(lp, f'  CAPTION {number}  ·  ', bold=True, size=8, color=RED)
    run(lp, label.upper(), bold=True, size=8, color=LIGHT)

    # card table
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, 'FFF7E9')
    set_top_border(cell, platform_hex, sz=20)

    # body text — line by line
    lines = text.strip().split('\n')
    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)

        if line == '':
            p.paragraph_format.space_after = Pt(4)
            continue

        # hashtags
        if line.startswith('#'):
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            run(p, line, size=9, color=platform_rgb)

        # placeholders
        elif '[' in line:
            p.paragraph_format.space_after = Pt(2)
            # split on [ and ] to color placeholders red
            import re
            parts = re.split(r'(\[.*?\])', line)
            for part in parts:
                if part.startswith('['):
                    run(p, part, size=10.5, color=RED, bold=True)
                else:
                    run(p, part, size=10.5, color=BODY)

        # quote line
        elif line.startswith(('"', '\u201c')):
            p.paragraph_format.space_after = Pt(2)
            run(p, line, italic=True, size=10.5, color=MID)

        else:
            p.paragraph_format.space_after = Pt(2)
            run(p, line, size=10.5, color=BODY)

    # copy instruction
    cp = cell.add_paragraph()
    cp.paragraph_format.space_before = Pt(10)
    cp.paragraph_format.space_after  = Pt(6)
    run(cp, '⌨  To copy: ', bold=True, size=8.5, color=LIGHT)
    run(cp, 'Click inside this box → Ctrl+A (Select All) → Ctrl+C (Copy)',
        size=8.5, color=LIGHT, italic=True)

    doc.add_paragraph()


# ── Platform section header ───────────────────────────────────
def platform_header(doc, platform, icon, desc, hex_color, rgb_color):
    # colored bar
    bar_tbl  = doc.add_table(rows=1, cols=1)
    bar_cell = bar_tbl.cell(0, 0)
    set_cell_bg(bar_cell, hex_color)
    bp = bar_cell.paragraphs[0]
    bp.paragraph_format.space_before = Pt(3)
    bp.paragraph_format.space_after  = Pt(3)
    run(bp, f'  {icon}  {platform.upper()}',
        bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))

    # description
    dp = doc.add_paragraph()
    dp.paragraph_format.space_before = Pt(6)
    dp.paragraph_format.space_after  = Pt(14)
    run(dp, desc, size=9.5, color=MID, italic=True)


# ══════════════════════════════════════════════════════════════
#  BUILD
# ══════════════════════════════════════════════════════════════
doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin   = Cm(2.8)
    sec.right_margin  = Cm(2.8)

# ── Red top bar ───────────────────────────────────────────────
top_tbl  = doc.add_table(rows=1, cols=1)
top_cell = top_tbl.cell(0, 0)
set_cell_bg(top_cell, 'EC1C2D')
tp = top_cell.paragraphs[0]
tp.paragraph_format.space_before = Pt(0)
tp.paragraph_format.space_after  = Pt(0)
run(tp, '  MIRACLE OF MIND  ·  VOLUNTEER HUB',
    bold=True, size=8, color=RGBColor(0xFF, 0xFF, 0xFF))

# ── Title ─────────────────────────────────────────────────────
doc.add_paragraph()
h1 = doc.add_paragraph()
h1.paragraph_format.space_before = Pt(4)
h1.paragraph_format.space_after  = Pt(2)
run(h1, 'Social Media Captions', bold=True, size=26, color=DARK)

sub = doc.add_paragraph()
sub.paragraph_format.space_before = Pt(0)
sub.paragraph_format.space_after  = Pt(4)
run(sub, 'Session Promotion Kit', size=13, color=MID)

# thin rule
rule_tbl  = doc.add_table(rows=1, cols=1)
rule_cell = rule_tbl.cell(0, 0)
set_cell_bg(rule_cell, 'E8DECE')
rp = rule_cell.paragraphs[0]
rp.paragraph_format.space_before = Pt(0)
rp.paragraph_format.space_after  = Pt(0)
rp.add_run(' ')

# Intro note
doc.add_paragraph()
note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(0)
note.paragraph_format.space_after  = Pt(16)
run(note, '10 ready-to-post captions across 4 platforms. ', bold=True, size=10, color=DARK)
run(note, 'Fields in red must be filled in before posting. '
          'Use the Volunteer Hub website for one-click copying.',
    size=10, color=MID)

# ══════════════════════════════════════════════════════════════
#  INSTAGRAM
# ══════════════════════════════════════════════════════════════
platform_header(doc, 'Instagram', '📸', '4 captions — quote posts, session invites, benefit-led, and social proof.',
                '833AB4', PLATFORM_COLORS['Instagram'][1])

caption_card(doc, 1, 'Quote Post',
    platform_hex=PLATFORM_COLORS['Instagram'][0],
    platform_rgb=PLATFORM_COLORS['Instagram'][1],
    text='''\
"If your mind becomes a conscious process, it becomes the greatest miracle in existence." — Sadhguru

A free 7-minute meditation. Designed by Sadhguru. Backed by science. Available right now.
👉 Download Miracle of Mind — link in bio.

#MiracleOfMind #Sadhguru #Meditation #MentalWellbeing #IshaFoundation''')

caption_card(doc, 2, 'Session Invite',
    platform_hex=PLATFORM_COLORS['Instagram'][0],
    platform_rgb=PLATFORM_COLORS['Instagram'][1],
    text='''\
We're hosting a free 30-minute meditation session in your city — and you're invited. 🙏

No experience needed. Just show up.

📍 [Location]  ·  🗓️ [Date]  ·  ⏰ [Time]
Register: [link]

#MoMSession #FreeEvent #MiracleOfMind #IshaVolunteers''')

caption_card(doc, 3, 'Lifestyle / Benefit',
    platform_hex=PLATFORM_COLORS['Instagram'][0],
    platform_rgb=PLATFORM_COLORS['Instagram'][1],
    text='''\
7 minutes a day.
More clarity. Better sleep. Less stress.

That's what thousands of people are experiencing with the Miracle of Mind app by Sadhguru — and it's completely free.

Try it today 👇 [link]

#MiracleOfMind #7Minutes #MindfulLiving #Sadhguru''')

caption_card(doc, 4, 'Social Proof',
    platform_hex=PLATFORM_COLORS['Instagram'][0],
    platform_rgb=PLATFORM_COLORS['Instagram'][1],
    text='''\
1 million downloads in 15 hours. 🌍

People everywhere are discovering what 7 minutes of meditation can do. Join them — the Miracle of Mind app is free, guided by Sadhguru, and takes less time than your morning coffee.

📲 Download: [link]

#MiracleOfMind #Sadhguru #Meditation''')

# ══════════════════════════════════════════════════════════════
#  FACEBOOK
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
platform_header(doc, 'Facebook', '👥', '2 captions — community invite and personal story. Longer format works well here.',
                '1877F2', PLATFORM_COLORS['Facebook'][1])

caption_card(doc, 1, 'Community Invite',
    platform_hex=PLATFORM_COLORS['Facebook'][0],
    platform_rgb=PLATFORM_COLORS['Facebook'][1],
    text='''\
Hi friends! I'm volunteering with Isha Foundation to bring something special to our community — a free Miracle of Mind meditation session.

In just 30 minutes, you'll learn a simple 7-minute practice designed by Sadhguru that you can use every single day to reduce stress, improve focus, and feel more at ease.

📍 [Location]  ·  🗓️ [Date & Time]

No prior experience needed — just bring yourself. Feel free to share this with anyone who could use a little more calm in their life. 🙏

Register here: [link]''')

caption_card(doc, 2, 'Personal Story',
    platform_hex=PLATFORM_COLORS['Facebook'][0],
    platform_rgb=PLATFORM_COLORS['Facebook'][1],
    text='''\
A few months ago, I started doing a 7-minute meditation every morning. I was skeptical at first — 7 minutes didn't seem like nearly enough.

But something shifted. I started responding instead of reacting. Sleep got better. The mental chatter quieted down.

That practice is the Miracle of Mind — a free app by Sadhguru, available to anyone. I'm now volunteering to share it as widely as possible.

If you're curious, download it here: [link] 🌱''',
    tip='Make this your own — replace or add to the personal experience with something true to your own practice.')

# ══════════════════════════════════════════════════════════════
#  LINKEDIN
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
platform_header(doc, 'LinkedIn', '💼', '2 captions — workplace wellness and thought leadership. Keep tone professional.',
                '0A66C2', PLATFORM_COLORS['LinkedIn'][1])

caption_card(doc, 1, 'Workplace Wellness',
    platform_hex=PLATFORM_COLORS['LinkedIn'][0],
    platform_rgb=PLATFORM_COLORS['LinkedIn'][1],
    text='''\
Burnout. Anxiety. Lack of focus. These aren't personal failures — they're signals.

Isha Foundation is offering free Miracle of Mind Experience Sessions for organizations — a 30-minute session introducing employees to a science-backed 7-minute daily meditation designed by Sadhguru.

20+ studies from world-class institutions validate its effectiveness. Over 1 million people downloaded the app within 15 hours of launch.

If you're in HR, L&D, or lead a team and want to bring this to your organization, reach out or visit: [link]

#WorkplaceWellness #MentalHealth #MiracleOfMind #IshaFoundation''')

caption_card(doc, 2, 'Thought Leadership',
    platform_hex=PLATFORM_COLORS['LinkedIn'][0],
    platform_rgb=PLATFORM_COLORS['LinkedIn'][1],
    text='''\
We spend years building skills, strategies, and systems.
But the mind running all of it? Often left unattended.

The Miracle of Mind app by Sadhguru offers a simple 7-minute daily practice — free, guided, and validated by science — to help you take charge of your mental wellbeing.

I've found it genuinely useful. Sharing it here in case it helps someone else too.
📲 [link]

#MiracleOfMind #Sadhguru #Leadership #Wellbeing''')

# ══════════════════════════════════════════════════════════════
#  X / TWITTER
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
platform_header(doc, 'X  (Twitter)', '✦', '2 short posts — punchy and under 280 characters. No edits needed, just replace the link.',
                '1C1C1C', PLATFORM_COLORS['X'][1])

caption_card(doc, 1, 'Punchy Stat',
    platform_hex='333333',
    platform_rgb=PLATFORM_COLORS['X'][1],
    text='''\
7 minutes of meditation.
50% reduction in stress.
Free app by Sadhguru.

That's Miracle of Mind. → [link]''')

caption_card(doc, 2, 'Quote',
    platform_hex='333333',
    platform_rgb=PLATFORM_COLORS['X'][1],
    text='''\
"If your mind becomes a conscious process, it becomes the greatest miracle in existence."
— Sadhguru

Free 7-min daily meditation: [link] #MiracleOfMind''')

# ── Footer ────────────────────────────────────────────────────
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(10)
run(fp, 'All content © Isha Foundation  ·  Miracle of Mind Volunteer Hub  ·  ',
    size=8, color=LIGHT, italic=True)
run(fp, 'aravindmurari.github.io/MoM-Isha',
    size=8, color=RED, italic=True)

# ── Save ──────────────────────────────────────────────────────
out = '/Users/aravindmurari/Downloads/MoM_Social_Media_Captions.docx'
doc.save(out)
print(f'Saved → {out}')
