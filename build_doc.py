from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── MoM brand colors ──────────────────────────────────────────
RED   = RGBColor(0xEC, 0x1C, 0x2D)   # MoM primary red
DARK  = RGBColor(0x1A, 0x1A, 0x1A)   # heading / label text
BODY  = RGBColor(0x33, 0x2E, 0x2A)   # body copy
MID   = RGBColor(0x6E, 0x66, 0x5B)   # secondary / muted
LIGHT = RGBColor(0xA0, 0x98, 0x90)   # notes, rules

# ── Helpers ───────────────────────────────────────────────────
def set_cell_bg(cell, hex6):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6)
    tcPr.append(shd)

def set_cell_borders(cell, sides, color, sz=6, val='single'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcB  = OxmlElement('w:tcBorders')
    all_sides = ('top', 'left', 'bottom', 'right')
    for s in all_sides:
        el = OxmlElement(f'w:{s}')
        if s in sides:
            el.set(qn('w:val'),   val)
            el.set(qn('w:sz'),    str(sz))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color)
        else:
            el.set(qn('w:val'),   'none')
            el.set(qn('w:sz'),    '0')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), 'auto')
        tcB.append(el)
    tcPr.append(tcB)

def run(para, text, bold=False, italic=False, size=None, color=None, underline=False):
    r = para.add_run(text)
    r.bold      = bold
    r.italic    = italic
    r.underline = underline
    if size:  r.font.size      = Pt(size)
    if color: r.font.color.rgb = color
    return r

def para(doc, text='', align=None, sb=0, sa=0):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        p.add_run(text)
    return p


# ══════════════════════════════════════════════════════════════
#  TEMPLATE RENDERER
#  Wraps one email template in a clean, branded card
# ══════════════════════════════════════════════════════════════
def email_card(doc, number, audience, subject, body_lines, tip):
    """
    Renders a single email template as a framed card:
      • Red top-border accent
      • Warm cream background
      • Numbered label + audience tag
      • Bold subject line
      • Body with smart styling per line type
      • Italicised volunteer tip at the bottom
    """

    # ── outer single-cell table ────────────────────────────────
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, 'FFF7E9')          # MoM warm cream

    # thick red top border, thin warm rule on other sides
    tcPr = cell._tc.get_or_add_tcPr()
    tcB  = OxmlElement('w:tcBorders')
    borders = {
        'top':    ('EC1C2D', '24'),
        'bottom': ('E8DECE', '4'),
        'left':   ('E8DECE', '4'),
        'right':  ('E8DECE', '4'),
    }
    for side, (color, sz) in borders.items():
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcB.append(el)
    tcPr.append(tcB)

    # ── label row: "EMAIL 1 OF 3  ·  AUDIENCE" ────────────────
    lp = cell.paragraphs[0]
    lp.paragraph_format.space_before = Pt(10)
    lp.paragraph_format.space_after  = Pt(2)
    run(lp, f'EMAIL {number} OF 3  ·  ', bold=True, size=8, color=RED)
    run(lp, audience.upper(), bold=True, size=8, color=LIGHT)

    # ── subject line ──────────────────────────────────────────
    sp = cell.add_paragraph()
    sp.paragraph_format.space_before = Pt(6)
    sp.paragraph_format.space_after  = Pt(2)
    run(sp, 'Subject: ', bold=True, size=11, color=DARK)
    run(sp, subject,     bold=False, size=11, color=DARK)

    # thin divider rule
    dp = cell.add_paragraph()
    dp.paragraph_format.space_before = Pt(4)
    dp.paragraph_format.space_after  = Pt(8)
    r = dp.add_run('─' * 80)
    r.font.size      = Pt(6)
    r.font.color.rgb = RGBColor(0xCC, 0xC4, 0xBA)

    # ── body ──────────────────────────────────────────────────
    for line in body_lines:
        bp = cell.add_paragraph()
        bp.paragraph_format.space_before = Pt(0)

        if line == '':
            bp.paragraph_format.space_after = Pt(5)
            continue

        # bullet points
        if line.startswith('•'):
            bp.paragraph_format.left_indent  = Inches(0.2)
            bp.paragraph_format.space_after  = Pt(2)
            run(bp, line, size=10.5, color=BODY)

        # testimonial quotes
        elif line.startswith(('"', '\u201c')):
            bp.paragraph_format.left_indent  = Inches(0.2)
            bp.paragraph_format.space_after  = Pt(2)
            run(bp, line, italic=True, size=10.5, color=MID)

        # attribution lines
        elif line.startswith('–') or line.startswith('—'):
            bp.paragraph_format.left_indent  = Inches(0.3)
            bp.paragraph_format.space_after  = Pt(6)
            run(bp, line, italic=True, size=9.5, color=LIGHT)

        # sign-off
        elif any(line.startswith(x) for x in ('In Love', 'Warm Regards', '[Your Name]')):
            bp.paragraph_format.space_after = Pt(2)
            run(bp, line, italic=True, size=10.5, color=MID)

        # sub-headings within the body (e.g. "Why Offer This Session?")
        elif line.endswith('?') and len(line) < 60 and ',' not in line:
            bp.paragraph_format.space_before = Pt(6)
            bp.paragraph_format.space_after  = Pt(3)
            run(bp, line, bold=True, size=10.5, color=DARK)

        # placeholder fields
        elif line.startswith('[') or line.startswith('Where:') or line.startswith('When:'):
            bp.paragraph_format.space_after = Pt(2)
            run(bp, line, size=10.5, color=RED)

        else:
            bp.paragraph_format.space_after = Pt(3)
            run(bp, line, size=10.5, color=BODY)

    # ── volunteer tip ─────────────────────────────────────────
    tp = cell.add_paragraph()
    tp.paragraph_format.space_before = Pt(10)
    tp.paragraph_format.space_after  = Pt(10)
    run(tp, '✦ Volunteer tip:  ', bold=True, size=9, color=RED)
    run(tp, tip, italic=True, size=9, color=MID)


# ══════════════════════════════════════════════════════════════
#  BUILD
# ══════════════════════════════════════════════════════════════
doc = Document()

# Page margins
for sec in doc.sections:
    sec.top_margin    = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin   = Cm(2.8)
    sec.right_margin  = Cm(2.8)

# ── HEADER ────────────────────────────────────────────────────
# Red top bar
top_tbl  = doc.add_table(rows=1, cols=1)
top_cell = top_tbl.cell(0, 0)
set_cell_bg(top_cell, 'EC1C2D')
tp = top_cell.paragraphs[0]
tp.paragraph_format.space_before = Pt(0)
tp.paragraph_format.space_after  = Pt(0)
run(tp, '  MIRACLE OF MIND  ·  VOLUNTEER HUB',
    bold=True, size=8, color=RGBColor(0xFF,0xFF,0xFF))

# Title block
doc.add_paragraph()
h1 = doc.add_paragraph()
h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
h1.paragraph_format.space_before = Pt(4)
h1.paragraph_format.space_after  = Pt(2)
run(h1, 'Email Templates', bold=True, size=26, color=DARK)

sub = doc.add_paragraph()
sub.paragraph_format.space_before = Pt(0)
sub.paragraph_format.space_after  = Pt(4)
run(sub, 'Session Promotion Kit', size=13, color=MID)

# Thin warm rule under title
rule_tbl  = doc.add_table(rows=1, cols=1)
rule_cell = rule_tbl.cell(0, 0)
set_cell_bg(rule_cell, 'E8DECE')
rp = rule_cell.paragraphs[0]
rp.paragraph_format.space_before = Pt(0)
rp.paragraph_format.space_after  = Pt(0)
rp.add_run(' ')

# Intro note
doc.add_paragraph()
note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(0)
note.paragraph_format.space_after  = Pt(16)
run(note, 'Three templates for different audiences. ', bold=True, size=10, color=DARK)
run(note,
    'Replace all fields in red — [ brackets ] — with your details before sending.',
    size=10, color=MID)

# ── EMAIL TEMPLATES ───────────────────────────────────────────

email_card(doc,
    number   = 1,
    audience = 'Corporate & Organizational Outreach',
    subject  = 'Free Offering — Miracle of Mind Meditation EXPERIENCE Session for Your Organization',
    body_lines = [
        'Greetings!',
        '',
        'In today\'s fast-paced world, stress, anxiety, and burnout have become common challenges across workplaces and educational institutions. At Isha Foundation, we are offering a simple yet transformative solution – the Miracle of Mind Experience Session.',
        '',
        'This is a 30-minute in-person session designed to introduce participants to Miracle of Mind, a powerful 7-minute guided meditation by Sadhguru. The practice is accessible to all, requires no prior experience, and can be easily integrated into one\'s daily routine.',
        '',
        'Why Offer This Session?',
        '• Enhance mental clarity, focus, and emotional balance',
        '• Reduce stress and anxiety',
        '• Improve overall wellbeing and productivity',
        '• Cultivate a positive and energized atmosphere among employees, staff and students',
        '',
        'The Miracle of Mind app clocked over one million downloads within 15 hours of launch. Here\'s what some participants are saying:',
        '',
        '\u201cAfter a week of meditation, everything changed. I now handle tasks with ease, experience less friction, and find genuine joy in my work.\u201d',
        '– Natesh Bhat, Engineer, Tumkur, India',
        '',
        '\u201cI found myself feeling calm in previously stressful situations and less identified with my thoughts and emotions.\u201d',
        '– Tom Davis, Software Developer, North Carolina, USA',
        '',
        'How to Get Started',
        '',
        'If your institution or company is interested in hosting this session, please email us at: corporate@ishausa.org',
        '',
        'Our team will connect with you to schedule the session at a time that suits you best.',
        '',
        'This is a unique opportunity to introduce a scientifically backed, time-efficient wellbeing practice to your community.',
        '',
        'We look forward to partnering with you.',
        '',
        'In Love, Light, Laughter',
        'Isha Volunteers',
    ],
    tip = 'Best sent to HR managers, department heads, or wellness coordinators. Follow up with a call if no response within 5 days.'
)

doc.add_paragraph()

email_card(doc,
    number   = 2,
    audience = 'Personal Invite — Friends, Family & Neighbors',
    subject  = "You're invited! Miracle of Mind Meditation EXPERIENCE Session",
    body_lines = [
        'Hi [Name / Friends / Neighbors],',
        '',
        'Namaskaram!',
        '',
        'I hope you\'re doing well. I wanted to personally invite you to a simple but powerful 30-minute meditation session that can make a real difference to your mental well-being.',
        '',
        'I\'ll be offering a Miracle of Mind (MoM) session — a guided 7-minute meditation designed by Sadhguru to bring clarity, calm, and focus to your everyday life.',
        '',
        'Where:   [Insert location]',
        'When:    [Insert date & time]',
        '',
        'What do you need? Just your presence. No prior experience needed!',
        '',
        'At the end of the session, you\'ll be able to download the free Miracle of Mind app and continue the practice at your convenience — anytime, anywhere.',
        '',
        'I would love for you to join. Feel free to bring a friend or family member too!',
        '',
        'Let me know if you have any questions.',
        '',
        'Warm Regards,',
        '[Your Name]',
    ],
    tip = 'Add a line or two about how the practice has personally benefited you — it makes this feel genuine rather than a broadcast.'
)

doc.add_paragraph()

email_card(doc,
    number   = 3,
    audience = 'Post-Session Follow-up',
    subject  = 'Thank you for joining — your Miracle of Mind practice',
    body_lines = [
        'Dear Participant,',
        '',
        'Thank you for joining us for today\'s Miracle of Mind Experience session. It was wonderful to have you with us, and we hope you continue to make use of the practice.',
        '',
        'For maximum benefit, the practice is recommended on a daily basis.',
        'Download the Miracle of Mind app here: isha.co/app',
        '',
        'To stay informed about future programs and Sadhguru\'s events, please fill out this form: isha.co/idystayconnected',
        '',
        'Wishing you continued wellbeing. Stay safe and healthy.',
        '',
        'In Love, Light, Laughter',
        'Isha Volunteers',
    ],
    tip = 'Send within 24 hours of the session while the experience is still fresh.'
)

# ── FOOTER ────────────────────────────────────────────────────
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(10)
run(fp, 'All content © Isha Foundation  ·  Miracle of Mind Volunteer Hub',
    size=8, color=LIGHT, italic=True)

# ── SAVE ──────────────────────────────────────────────────────
out = '/Users/aravindmurari/Downloads/MoM_Email_Templates.docx'
doc.save(out)
print(f'Saved → {out}')
