from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── MoM brand colors ──────────────────────────────────────────
RED     = RGBColor(0xEC, 0x1C, 0x2D)
TEAL    = RGBColor(0x7C, 0xB8, 0xA8)   # session-kit accent
DARK    = RGBColor(0x1A, 0x1A, 0x1A)
BODY    = RGBColor(0x33, 0x2E, 0x2A)
MID     = RGBColor(0x6E, 0x66, 0x5B)
LIGHT   = RGBColor(0xA0, 0x98, 0x90)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ───────────────────────────────────────────────────
def set_cell_bg(cell, hex6):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6)
    tcPr.append(shd)

def set_borders(cell, top_color, top_sz=20, other_color='E8DECE', other_sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    tcB  = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(top_sz) if side == 'top' else str(other_sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), top_color if side == 'top' else other_color)
        tcB.append(el)
    tcPr.append(tcB)

def run(para, text, bold=False, italic=False, size=None, color=None):
    r = para.add_run(text)
    r.bold   = bold
    r.italic = italic
    if size:  r.font.size      = Pt(size)
    if color: r.font.color.rgb = color
    return r

def spacer(doc, pts=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)

def section_header(doc, title, subtitle=None):
    """Teal-bar section divider."""
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, '7CB8A8')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run(p, f'  {title}', bold=True, size=11, color=WHITE)
    spacer(doc, 6)
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(12)
        run(sp, subtitle, size=9.5, color=MID, italic=True)

def talking_card(doc, label, lines, tip=None):
    """A single talking point card."""
    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after  = Pt(3)
    run(lp, f'  {label.upper()}', bold=True, size=8, color=TEAL)

    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, 'FFF7E9')
    set_borders(cell, '7CB8A8', top_sz=16)

    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)
        if line == '':
            p.paragraph_format.space_after = Pt(5)
            continue
        elif line.startswith('•'):
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(2)
            run(p, line, size=10.5, color=BODY)
        elif line.startswith('"') or line.startswith('\u201c'):
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(4)
            run(p, line, italic=True, size=10.5, color=MID)
        else:
            p.paragraph_format.space_after = Pt(3)
            run(p, line, size=10.5, color=BODY)

    if tip:
        tp = cell.add_paragraph()
        tp.paragraph_format.space_before = Pt(8)
        tp.paragraph_format.space_after  = Pt(6)
        run(tp, '✦ Tip:  ', bold=True, size=9, color=RED)
        run(tp, tip, italic=True, size=9, color=MID)
    else:
        ep = cell.add_paragraph()
        ep.paragraph_format.space_before = Pt(4)
        ep.paragraph_format.space_after  = Pt(4)
        ep.add_run('')

    spacer(doc, 10)

def objection_card(doc, objection, response):
    """Q&A objection card with red question, teal response."""
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, 'FFF7E9')
    set_borders(cell, 'EC1C2D', top_sz=16)

    qp = cell.paragraphs[0]
    qp.paragraph_format.space_before = Pt(6)
    qp.paragraph_format.space_after  = Pt(4)
    run(qp, objection, bold=True, size=10.5, color=DARK)

    ap = cell.add_paragraph()
    ap.paragraph_format.space_before = Pt(0)
    ap.paragraph_format.space_after  = Pt(8)
    ap.paragraph_format.left_indent  = Inches(0.15)
    run(ap, '↳  ', bold=True, size=10, color=TEAL)
    run(ap, response, italic=True, size=10.5, color=MID)

    spacer(doc, 8)

def script_line(cell, stage, text, is_stage=False, indent=False):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    if is_stage:
        p.paragraph_format.space_before = Pt(10)
        run(p, stage, bold=True, size=9, color=TEAL)
    elif stage:
        run(p, stage, bold=True, size=10.5, color=DARK)
        run(p, f'  {text}', size=10.5, color=BODY)
    else:
        run(p, text, italic=True, size=10.5, color=MID)


# ══════════════════════════════════════════════════════════════
#  BUILD
# ══════════════════════════════════════════════════════════════
doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin   = Cm(2.8)
    sec.right_margin  = Cm(2.8)

# ── Red top bar ───────────────────────────────────────────────
top_tbl  = doc.add_table(rows=1, cols=1)
top_cell = top_tbl.cell(0, 0)
set_cell_bg(top_cell, 'EC1C2D')
tp = top_cell.paragraphs[0]
tp.paragraph_format.space_before = Pt(0)
tp.paragraph_format.space_after  = Pt(0)
run(tp, '  MIRACLE OF MIND  ·  VOLUNTEER HUB',
    bold=True, size=8, color=WHITE)

# ── Title ─────────────────────────────────────────────────────
doc.add_paragraph()
h1 = doc.add_paragraph()
h1.paragraph_format.space_before = Pt(4)
h1.paragraph_format.space_after  = Pt(2)
run(h1, 'Talking Points', bold=True, size=26, color=DARK)

sub = doc.add_paragraph()
sub.paragraph_format.space_before = Pt(0)
sub.paragraph_format.space_after  = Pt(4)
run(sub, 'Session Promotion Kit', size=13, color=MID)

rule_tbl  = doc.add_table(rows=1, cols=1)
rule_cell = rule_tbl.cell(0, 0)
set_cell_bg(rule_cell, 'E8DECE')
rp = rule_cell.paragraphs[0]
rp.paragraph_format.space_before = Pt(0)
rp.paragraph_format.space_after  = Pt(0)
rp.add_run(' ')

doc.add_paragraph()
note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(0)
note.paragraph_format.space_after  = Pt(16)
run(note, 'Three sections: ', bold=True, size=10, color=DARK)
run(note, 'Key talking points to introduce MoM, responses to common objections, '
          'and a 2-minute conversation script.',
    size=10, color=MID)

# ══════════════════════════════════════════════════════════════
#  SECTION 1 — KEY TALKING POINTS
# ══════════════════════════════════════════════════════════════
section_header(doc, '1.  Key Talking Points',
    'Use these in any order depending on the conversation. Pick 2–3 that feel natural.')

talking_card(doc, 'Opening — how to bring it up',
    lines=[
        '"Have you heard about the Miracle of Mind app? It\'s a free meditation by Sadhguru — takes 7 minutes a day."',
        '',
        '"I\'ve been doing this 7-minute meditation every morning and it\'s genuinely made a difference. Have you tried anything like that?"',
        '',
        '"We\'re hosting a free 30-minute session nearby — no prior experience needed. Would you be open to trying something like that?"',
    ],
    tip='Lead with curiosity, not a pitch. A question lands better than a statement.')

talking_card(doc, 'The one-line pitch',
    lines=[
        '"It\'s a free app by Sadhguru — a guided 7-minute meditation you do once a day, backed by 20+ scientific studies."',
        '',
        '"Think of it as a gym for your mind — 7 minutes a day, completely free, works for anyone."',
    ])

talking_card(doc, 'Key benefits — pick 2–3 based on the person',
    lines=[
        '• Reduces stress and anxiety',
        '• Better sleep and more energy',
        '• Sharper focus and clarity at work',
        '• Helps you respond to situations rather than react',
        '• No experience, no equipment, no cost',
    ],
    tip='Listen first. If they mention sleep, lead with sleep. If they mention work stress, lead with focus.')

talking_card(doc, 'Social proof — facts to cite with confidence',
    lines=[
        '"The app hit 1 million downloads in just 15 hours of launch."',
        '',
        '"Over 20 studies from world-class institutions validate the practices."',
        '',
        '"7 minutes of this meditation measurably increases alpha waves in the brain — that\'s physical relaxation, not just a feeling."',
    ])

talking_card(doc, 'The ask — how to close',
    lines=[
        '"Would you be open to trying it for just 7 days? That\'s all it takes to feel a difference."',
        '',
        '"Can I send you the link? It takes 2 minutes to download."',
        '',
        '"We have a free session coming up — would you want to come and try it in person before committing to anything?"',
    ],
    tip='Always offer two options: download the app now, or attend a session. Gives them a low-pressure choice.')

# ══════════════════════════════════════════════════════════════
#  SECTION 2 — HANDLING OBJECTIONS
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
section_header(doc, '2.  Handling Common Objections',
    'Stay warm and curious — never argue. Acknowledge, then reframe.')

objection_card(doc,
    '"I don\'t have time."',
    '"It\'s 7 minutes — less than a coffee break. Most people fit it in right after waking up."')

objection_card(doc,
    '"I\'ve tried meditation before and it didn\'t work for me."',
    '"This is specifically designed for people who struggle with traditional meditation — it\'s guided, short, and structured. Very different from just sitting and trying to clear your mind."')

objection_card(doc,
    '"Is this religious?"',
    '"Not at all — it\'s a wellbeing tool, like yoga. People from all backgrounds and beliefs use it. The science stands on its own."')

objection_card(doc,
    '"Is it really free? What\'s the catch?"',
    '"Completely free. No subscription, no credit card. It\'s Isha Foundation\'s offering to the world."')

objection_card(doc,
    '"I\'m already doing okay mentally."',
    '"Totally fair — most people who use it aren\'t in crisis. They find it sharpens focus and adds a sense of ease to days that are already going well."')

# ══════════════════════════════════════════════════════════════
#  SECTION 3 — 2-MINUTE CONVERSATION SCRIPT
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
section_header(doc, '3.  2-Minute Conversation Script',
    'A natural flow from opener to close. Adapt it — don\'t read it word for word.')

tbl  = doc.add_table(rows=1, cols=1)
cell = tbl.cell(0, 0)
set_cell_bg(cell, 'FFF7E9')
set_borders(cell, '7CB8A8', top_sz=20)

# Stage 1
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(2)
run(p, 'OPEN  ·  Gauge interest', bold=True, size=8.5, color=TEAL)

lines_script = [
    (None, '"Hey [Name], have you heard about the Miracle of Mind app? It\'s a free meditation by Sadhguru — takes 7 minutes a day."', False),
    ('', '', False),
    ('If yes →', '"Oh great! Have you tried it yet?"', True),
    ('If no  →', '"It\'s a guided meditation, completely free, 7 minutes. Over a million people downloaded it in the first 15 hours."', True),
    ('STAGE', 'MAKE IT PERSONAL', False),
    (None, '"I know you\'ve been [stressed at work / not sleeping well / dealing with a lot lately] — this is exactly the kind of thing that helps with that. And it\'s not just feel-good stuff — there are 20+ scientific studies behind it."', False),
    ('STAGE', 'ADD YOUR OWN EXPERIENCE', False),
    (None, '"I\'ve been doing it myself and honestly, I notice a difference on days I skip it. It\'s one of those things that sounds too simple to work, but it does."', False),
    ('STAGE', 'THE ASK', False),
    (None, '"Would you be open to downloading it and trying it for 7 days? Or if you\'d rather experience it in person first, we\'re hosting a free 30-minute session on [date]."', False),
    ('STAGE', 'CLOSE', False),
    (None, '"Here\'s the link — completely free, no account needed to start. Let me know what you think!"', False),
]

for label, text, indent in lines_script:
    if label == 'STAGE':
        sp = cell.add_paragraph()
        sp.paragraph_format.space_before = Pt(12)
        sp.paragraph_format.space_after  = Pt(4)
        run(sp, text, bold=True, size=8.5, color=TEAL)
    elif label == '':
        ep = cell.add_paragraph()
        ep.paragraph_format.space_before = Pt(0)
        ep.paragraph_format.space_after  = Pt(4)
    elif label:
        lp = cell.add_paragraph()
        lp.paragraph_format.space_before = Pt(2)
        lp.paragraph_format.space_after  = Pt(2)
        lp.paragraph_format.left_indent  = Inches(0.2)
        run(lp, f'{label}  ', bold=True, size=9.5, color=DARK)
        run(lp, text, italic=True, size=9.5, color=MID)
    else:
        bp = cell.add_paragraph()
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after  = Pt(4)
        run(bp, text, italic=True, size=10.5, color=MID)

# closing note inside card
np = cell.add_paragraph()
np.paragraph_format.space_before = Pt(12)
np.paragraph_format.space_after  = Pt(8)
run(np, '✦ Remember:  ', bold=True, size=9, color=RED)
run(np, 'The script is a guide, not a script. The most effective conversations are genuine ones — '
        'share what the practice has meant to you personally.',
    italic=True, size=9, color=MID)

# ── Footer ────────────────────────────────────────────────────
spacer(doc, 16)
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(fp, 'All content © Isha Foundation  ·  Miracle of Mind Volunteer Hub  ·  ',
    size=8, color=LIGHT, italic=True)
run(fp, 'aravindmurari.github.io/MoM-Isha',
    size=8, color=RED, italic=True)

# ── Save ──────────────────────────────────────────────────────
out = '/Users/aravindmurari/Downloads/MoM_Talking_Points.docx'
doc.save(out)
print(f'Saved → {out}')
