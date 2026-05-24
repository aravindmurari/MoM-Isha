from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ── Helpers ───────────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15 if level == 1 else 12)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    else:
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p

def subheading(text):
    return heading(text, level=2)

def body(text, italic=False, bold=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.italic = italic
    run.bold   = bold
    if color:
        run.font.color.rgb = color
    return p

def divider():
    p = doc.add_paragraph('─' * 60)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    run = p.runs[0]
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


# ══════════════════════════════════════════════════════════════
# PART 1 — PROMPT FROM USER
# ══════════════════════════════════════════════════════════════

heading('User Prompt')
divider()

prompt_text = (
    "For the walkthrough video script, use a top-down approach — start with an overview of the intent "
    "and all three primary sections: Outreach Resources, Session Resources, and Take It Further (the third section).\n\n"
    "Make sure to communicate that the Session Resources section is accessible only after entering a passcode. "
    "Volunteers should have received this code via email upon completing their Boom Buddy training, and they "
    "need to enter it to unlock that section.\n\n"
    "Structure the script by walking through each section in order. The entire walkthrough should be no more "
    "than four to four and a half minutes. Keep it short and high-level — much of the site is self-evident, "
    "so there is no need to hand-hold every detail. The goal is simply to orient volunteers so they know "
    "how to navigate the site and can explore the details on their own.\n\n"
    "Additional requirements:\n"
    "- Start with Namaskaram.\n"
    "- In the Session Resources section, mention that volunteers can revisit the Boom Buddy training video "
    "at any time as a refresher.\n"
    "- At the close, let volunteers know that contact information for questions or additional support is "
    "also available at the bottom of the page.\n"
    "- Output the final script as a Word document that includes both this prompt and the script together."
)

body(prompt_text)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PART 2 — VIDEO SCRIPT
# ══════════════════════════════════════════════════════════════

heading('Boom Buddy Resource Kit — Walkthrough Video Script')
body('Target runtime: ~4 minutes', italic=True, color=RGBColor(0x88, 0x88, 0x88))
divider()

# INTRO
subheading('[INTRO — 0:00]')
body(
    "Namaskaram, and welcome to the Boom Buddy Resource Kit for Miracle of Mind. "
    "This is your central hub as a volunteer — everything you need for outreach and facilitating sessions,"
    "all in one place. This walkthrough will give you a quick overview of the three sections so you can get started right away."
)

# OVERVIEW
subheading('[OVERVIEW — 0:20]')
body(
    "The page is organized top to bottom into three sections: Outreach Resources, Session Resources, "
    "and Take It Further. At the top nav bar you can jump directly to any section, and you can always "
    "come back to this walkthrough video using the Watch Walkthrough button."
)

# SECTION 1
subheading('[SECTION 1: OUTREACH RESOURCES — 0:40]')
body(
    "The first section is Outreach Resources — everything you need to get people to a session."
)
body(
    "At the top you have an App Features Guide — a walkthrough of the Miracle of Mind app so you can "
    "confidently answer questions during outreach. Next is Booth Setup, with a guide and video for setting "
    "up your session venue."
)
body(
    "Below that are Flyer Templates — twelve ready-made designs. Use the Flyer Generator to fill in "
    "your event details and download a print-ready PNG, or open in Canva for full customization."
)
body(
    "Then expandable cards for Email Templates, WhatsApp Templates, and Social Media Captions — "
    "each organized by audience type: friends and family, universities, corporations, nonprofits, and more. "
    "Click to expand, pick your audience, personalize the text, and you're ready to send."
)
body(
    "You'll also find Printables for your session venue, and Other Useful Content with reference "
    "materials like the research summary, a corporate presentation deck, and verified Sadhguru quotes."
)

# SECTION 2
subheading('[SECTION 2: SESSION RESOURCES — 2:05]')
body(
    "The second section is Session Resources — and this is for trained Boom Buddy volunteers only."
)
body(
    "When you arrive here, you'll see a passcode prompt. When you completed your Boom Buddy training, "
    "you received an email with your access code. Enter that code here to unlock the full session kit. "
    "Once entered, the site remembers you on that device so you won't need to re-enter it each time."
)
body(
    "And if you ever want to go through the training again as a refresher, there's an option right on "
    "this page to revisit the Boom Buddy training video."
)
body(
    "Once you're inside, you'll find three things. First, the Session Script — with tabs for the Full "
    "Session Script and an FAQ. These are your dialogue guides for running the session. "
    "Second, a Session Flow Checklist — step-by-step with timing cues for each step; check things off "
    "as you go. And third, the Session Media player — all the videos for the session are right here. "
    "Click any item from the playlist to load and play it. When you're ready to present, hit Open "
    "Presentation View to send the video full-screen to your projector or second monitor — while that "
    "window is open, playback is locked to the presentation screen only. "
    "If you're running a session without reliable internet, you can also download the videos for offline use ahead of time."
)

# SECTION 3
subheading('[SECTION 3: TAKE IT FURTHER — 3:20]')
body(
    "At the bottom is Take It Further — four cards for volunteers who want to expand their involvement."
)
body(
    "Host a Group Session if you want to bring Miracle of Mind or other Isha programs to a workplace "
    "or community. MoM for Healthcare Professionals — a CME version is coming soon. Explore a Campus "
    "Club if you're interested in bringing MoM to your university. And Get in Touch for any general "
    "questions or support."
)
body(
    "Contact details for all of these are right there on the cards — and if you need anything else, "
    "the full contact information is also available at the bottom of the page."
)

# CLOSE
subheading('[CLOSE — 3:50]')
body(
    "That's it. Outreach Resources to spread the word, Session Resources to run the session, and "
    "Take It Further to deepen your involvement. Everything on the site is self-explanatory once "
    "you're in — this was just to help you get oriented. If you have any questions or need additional "
    "support, the contact info is at the bottom of the page. Thank you for being a Boom Buddy — go make it happen."
)

divider()
body('Estimated runtime: ~4:00', italic=True, color=RGBColor(0x88, 0x88, 0x88))

# ── Save ──────────────────────────────────────────────────────
out = '/Users/aravindmurari/Downloads/MoM_Walkthrough_Video_Script.docx'
doc.save(out)
print(f'Saved: {out}')
